# ================================
# Contract Generator Service - Version 2.0
# Professional International Contract Generator
# Tech: Python + FastAPI + SQLAlchemy + Jinja2
# Run: uvicorn main:app --host 0.0.0.0 --port 4000 --reload
# ================================

from fastapi import FastAPI, Form, Request, Depends, HTTPException
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from docx import Document
from datetime import datetime
from sqlalchemy.orm import Session
from typing import List
import os
import uuid
from num2words import num2words

# Local imports
from database import init_db, get_db, Contract
from pdf_generator import generate_pdf_contract


# ================================
# Configuration
# ================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "generated")
TEMPLATE_DIR = os.path.join(BASE_DIR, "templates")
CLAUSE_DIR = os.path.join(BASE_DIR, "clauses")

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Initialize database
init_db()

# FastAPI app
app = FastAPI(
    title="ContractPro - Professional Contract Generator",
    description="Generate legally compliant international service agreements",
    version="2.0"
)

# Static files and templates
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# Currency configuration
CURRENCIES = {
    "USD": {"symbol": "$", "name": "US Dollars", "rate": 1},
    "EUR": {"symbol": "€", "name": "Euros", "rate": 1.08},
    "GBP": {"symbol": "£", "name": "Pounds Sterling", "rate": 1.27},
    "AED": {"symbol": "AED ", "name": "UAE Dirhams", "rate": 0.27},
    "SAR": {"symbol": "SAR ", "name": "Saudi Riyals", "rate": 0.27},
    "BHD": {"symbol": "BHD ", "name": "Bahraini Dinars", "rate": 2.65},
    "QAR": {"symbol": "QAR ", "name": "Qatari Riyals", "rate": 0.27},
    "KWD": {"symbol": "KWD ", "name": "Kuwaiti Dinars", "rate": 3.25},
    "OMR": {"symbol": "OMR ", "name": "Omani Riyals", "rate": 2.60},
    "PKR": {"symbol": "PKR ", "name": "Pakistani Rupees", "rate": 0.0036}
}


# ================================
# Utility Functions
# ================================

def load_clause(name: str) -> str:
    """Load a clause template from file"""
    path = os.path.join(CLAUSE_DIR, f"{name}.txt")
    if not os.path.exists(path):
        return ""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_granular_service(service_id: str) -> str:
    """Load individual service file from category subdirectory
    
    Args:
        service_id: Format 'category_servicename' (e.g., 'it_app_development')
    
    Returns:
        Service content as string
    """
    # Split service_id into category and service name
    parts = service_id.split('_', 1)
    if len(parts) != 2:
        return ""
    
    category, service_name = parts
    path = os.path.join(CLAUSE_DIR, category, f"{service_name}.txt")
    
    if not os.path.exists(path):
        return ""
    
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def group_services_by_category(service_ids: list) -> dict:
    """Group service IDs by their category
    
    Args:
        service_ids: List of service IDs (e.g., ['it_app_development', 'finance_bookkeeping'])
    
    Returns:
        Dict with category names as keys and list of service IDs as values
    """
    grouped = {}
    category_names = {
        'it': 'IT & Software Services',
        'finance': 'Finance & Accounting Services',
        'hr': 'Human Resources & Payroll Services',
        'business': 'Business Consulting Services'
    }
    
    for service_id in service_ids:
        parts = service_id.split('_', 1)
        if len(parts) == 2:
            category = parts[0]
            if category not in grouped:
                grouped[category] = {
                    'name': category_names.get(category, category.title()),
                    'services': []
                }
            grouped[category]['services'].append(service_id)
    
    return grouped


def generate_docx_contract(data: dict) -> str:
    """Generate DOCX contract from template"""
    doc = Document(os.path.join(TEMPLATE_DIR, "master_contract.docx"))

    # Build services block with category grouping
    services_block = ""
    grouped_services = group_services_by_category(data["services"])
    

    for category, category_data in grouped_services.items():
        # Add category heading
        services_block += f"====> {category_data['name'].upper()}\n\n"
        
        # Add each service in this category
        for service_id in category_data['services']:
            service_content = load_granular_service(service_id)

            if service_content:
                services_block += service_content + "\n\n"
    


    # Load service provider and bank details
    service_provider_block = load_clause("service_provider")
    bank_details_block = load_clause("bank_details")
    
    # Replacements
    replacements = {
        "{{CLIENT_NAME}}": data["client_name"],
        "{{CLIENT_ADDRESS}}": data.get("client_address", ""),
        "{{COUNTRY}}": data["country"],
        "{{EFFECTIVE_DATE}}": data["date"],
        "{{FEES_AMOUNT}}": data["fees"],
        "{{FEES_IN_WORDS}}": data["fees_words"],
        "{{CURRENCY_SYMBOL}}": data["currency_symbol"],
        "{{CURRENCY_NAME}}": data["currency_name"],
        "{{USD_EQUIVALENT}}": data["usd_equivalent"],
        "{{CONTRACT_DURATION}}": data["contract_duration"],
        "{{SERVICES_BLOCK}}": services_block,
        "{{SERVICE_PROVIDER_BLOCK}}": service_provider_block,
        "{{BANK_DETAILS_BLOCK}}": bank_details_block
    }

    # Replace in paragraphs (preserve formatting by replacing in runs)
    services_replaced = False
    for p in doc.paragraphs:
        # Check if paragraph contains any placeholder
        if any(key in p.text for key in replacements.keys()):
            # Get the full paragraph text
            full_text = p.text
            
            # Replace all placeholders
            for key, value in replacements.items():
                if key in full_text:
                    if key == "{{SERVICES_BLOCK}}":
                        print(f"DEBUG: Found SERVICES_BLOCK placeholder in paragraph")
                        services_replaced = True
                    full_text = full_text.replace(key, value)
            
            # Clear existing runs and add new text
            # This preserves paragraph formatting but replaces content
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = full_text
            else:
                p.add_run(full_text)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Check if cell contains any placeholder
                if any(key in cell.text for key in replacements.keys()):
                    full_text = cell.text
                    
                    for key, value in replacements.items():
                        if key in full_text:
                            if key == "{{SERVICES_BLOCK}}":
                                print(f"DEBUG: Found SERVICES_BLOCK placeholder in table cell")
                                services_replaced = True
                            full_text = full_text.replace(key, value)
                    
                    # Replace cell text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text
                        else:
                            paragraph.add_run(full_text)
    


    # Save file
    file_id = str(uuid.uuid4())
    file_path = os.path.join(OUTPUT_DIR, f"contract_{file_id}.docx")
    doc.save(file_path)
    
    return file_path, file_id


def save_contract_to_db(data: dict, file_path: str, file_id: str, db: Session) -> Contract:
    """Save contract details to database"""
    contract = Contract(
        client_name=data["client_name"],
        client_address=data["client_address"],
        client_email=data.get("client_email", ""),
        client_phone=data.get("client_phone", ""),
        country=data["country"],
        fees=data["fees"],
        fees_numeric=float(data["fees"].replace(",", "")),
        fees_words=data["fees_words"],
        currency=data["currency"],
        currency_symbol=data["currency_symbol"],
        currency_name=data["currency_name"],
        usd_equivalent=data["usd_equivalent"],
        contract_duration=data["contract_duration"],
        services=", ".join([s.title() for s in data["services"]]),
        effective_date=data["date"],
        file_path=file_path,
        file_id=file_id
    )
    
    db.add(contract)
    db.commit()
    db.refresh(contract)
    
    return contract


# ================================
# Routes - Pages
# ================================

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    """Home page with contract generation form"""
    return templates.TemplateResponse(
        "pages/home.html",
        {"request": request, "active_page": "home"}
    )


@app.get("/contracts", response_class=HTMLResponse)
async def contracts_page(
    request: Request,
    page: int = 1,
    db: Session = Depends(get_db)
):
    """Contracts management page"""
    per_page = 20
    offset = (page - 1) * per_page
    
    # Get total count
    total = db.query(Contract).count()
    total_pages = (total + per_page - 1) // per_page
    
    # Get contracts for current page
    contracts = db.query(Contract)\
        .order_by(Contract.created_at.desc())\
        .offset(offset)\
        .limit(per_page)\
        .all()
    
    # Format dates for display
    for contract in contracts:
        contract.created_at = contract.created_at.strftime("%Y-%m-%d %H:%M")
    
    return templates.TemplateResponse(
        "pages/contracts.html",
        {
            "request": request,
            "active_page": "contracts",
            "contracts": contracts,
            "page": page,
            "total_pages": total_pages,
            "total": total
        }
    )


@app.get("/about", response_class=HTMLResponse)
async def about_page(request: Request):
    """About page"""
    return templates.TemplateResponse(
        "pages/about.html",
        {"request": request, "active_page": "about"}
    )


# ================================
# Routes - Contract Generation
# ================================

@app.post("/generate")
async def generate_contract(
    client_name: str = Form(...),
    client_address: str = Form(...),
    client_email: str = Form(...),
    client_phone: str = Form(...),
    country: str = Form(...),
    contract_start_date: str = Form(...),
    fees: str = Form(...),
    currency: str = Form("USD"),
    duration: str = Form("12 Months"),
    services: List[str] = Form([]),
    bilingual: str = Form(None),  # Checkbox value
    db: Session = Depends(get_db)
):
    """Generate contract and save to database"""
    
    # Validation
    if not services:
        raise HTTPException(status_code=400, detail="Please select at least one service")
    
    try:
        # Clean and parse fee amount
        amount = float(fees.replace(",", ""))
        if amount <= 0:
            raise ValueError("Amount must be positive")
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid fee amount")
    
    # Get currency data
    currency_data = CURRENCIES.get(currency)
    if not currency_data:
        raise HTTPException(status_code=400, detail="Invalid currency")
    
    # Calculate USD equivalent
    usd_equivalent = round(amount * currency_data["rate"], 2)
    
    # Convert amount to words
    try:
        fees_in_words = num2words(amount, to="currency", currency=currency).title()
    except:
        fees_in_words = f"{amount} {currency}"
    
    # Prepare data
    data = {
        "client_name": client_name,
        "client_address": client_address,
        "client_email": client_email,
        "client_phone": client_phone,
        "country": country,
        "fees": f"{amount:,.2f}",
        "fees_words": fees_in_words,
        "currency": currency,
        "currency_symbol": currency_data["symbol"],
        "currency_name": currency_data["name"],
        "usd_equivalent": f"{usd_equivalent:,.2f}",
        "contract_duration": duration,
        "date": contract_start_date,
        "services": services
    }
    
    # Generate DOCX
    file_path, file_id = generate_docx_contract(data)
    
    # Generate PDF
    pdf_file_path = file_path.replace('.docx', '.pdf')
    
    # Prepare services block for PDF with category grouping
    services_block = ""
    grouped_services = group_services_by_category(services)
    
    for category, category_data in grouped_services.items():
        # Add category heading
        services_block += f"====> {category_data['name'].upper()}\n\n"
        
        # Add each service in this category
        for service_id in category_data['services']:
            service_content = load_granular_service(service_id)
            if service_content:
                services_block += service_content + "\n\n"
    
    # Check if bilingual is requested
    is_bilingual = bilingual == "true" if bilingual else False
    
    # Add services_block to data for PDF
    data_for_pdf = {**data, "services_block": services_block}
    pdf_path = generate_pdf_contract(data_for_pdf, pdf_file_path, bilingual=is_bilingual)
    
    # Save to database
    contract = save_contract_to_db(data, file_path, file_id, db)
    
    # Return success page with download links
    return templates.TemplateResponse("pages/success.html", {
        "request": {},
        "client_name": client_name,
        "country": country,
        "fees": f"{amount:,.2f}",
        "currency_symbol": currency_data["symbol"],
        "currency_name": currency_data["name"],
        "duration": duration,
        "effective_date": contract_start_date,
        "bilingual": is_bilingual,
        "docx_url": f"/download/{contract.id}/docx",
        "pdf_url": f"/download/{contract.id}/pdf"
    })


# ================================
# Routes - Contract Management
# ================================

@app.get("/api/contract/{contract_id}")
async def get_contract(contract_id: int, db: Session = Depends(get_db)):
    """Get contract details by ID"""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    return {
        "id": contract.id,
        "client_name": contract.client_name,
        "country": contract.country,
        "fees": contract.fees,
        "currency": contract.currency,
        "currency_symbol": contract.currency_symbol,
        "currency_name": contract.currency_name,
        "usd_equivalent": contract.usd_equivalent,
        "duration": contract.contract_duration,
        "services": contract.services,
        "effective_date": contract.effective_date,
        "created_at": contract.created_at.strftime("%Y-%m-%d %H:%M")
    }


@app.get("/download/{contract_id}/{format}")
async def download_contract(
    contract_id: int,
    format: str,
    db: Session = Depends(get_db)
):
    """Download contract in specified format (docx or pdf)"""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    if format.lower() == "docx":
        # Return existing DOCX file
        if not os.path.exists(contract.file_path):
            raise HTTPException(status_code=404, detail="Contract file not found")
        
        return FileResponse(
            contract.file_path,
            filename=f"Service_Agreement_{contract.client_name.replace(' ', '_')}_{contract.effective_date.replace(' ', '_')}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    
    elif format.lower() == "pdf":
        # Generate PDF
        pdf_path = contract.file_path.replace('.docx', '.pdf')
        
        # Prepare data for PDF generation
        # Parse services from database (stored as comma-separated string)
        service_ids = [s.strip() for s in contract.services.split(", ")]
        
        # Build services block with category grouping
        services_block = ""
        grouped_services = group_services_by_category(service_ids)
        
        for category, category_data in grouped_services.items():
            # Add category heading
            services_block += f"====> {category_data['name'].upper()}\n\n"
            
            # Add each service in this category
            for service_id in category_data['services']:
                service_content = load_granular_service(service_id)
                if service_content:
                    services_block += service_content + "\n\n"
        
        pdf_data = {
            "client_name": contract.client_name,
            "client_address": contract.client_address or "N/A",
            "client_email": contract.client_email or "N/A",
            "client_phone": contract.client_phone or "N/A",
            "country": contract.country,
            "fees": contract.fees,
            "fees_words": contract.fees_words,
            "currency_symbol": contract.currency_symbol,
            "currency_name": contract.currency_name,
            "usd_equivalent": contract.usd_equivalent,
            "contract_duration": contract.contract_duration,
            "date": contract.effective_date,
            "services_block": services_block
        }
        
        generate_pdf_contract(pdf_data, pdf_path)
        
        return FileResponse(
            pdf_path,
            filename=f"Service_Agreement_{contract.client_name.replace(' ', '_')}_{contract.effective_date.replace(' ', '_')}.pdf",
            media_type="application/pdf"
        )
    
    else:
        raise HTTPException(status_code=400, detail="Invalid format. Use 'docx' or 'pdf'")


@app.delete("/api/contract/{contract_id}")
async def delete_contract(contract_id: int, db: Session = Depends(get_db)):
    """Delete a contract"""
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    
    if not contract:
        raise HTTPException(status_code=404, detail="Contract not found")
    
    # Delete files
    if os.path.exists(contract.file_path):
        os.remove(contract.file_path)
    
    pdf_path = contract.file_path.replace('.docx', '.pdf')
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    
    # Delete from database
    db.delete(contract)
    db.commit()
    
    return {"message": "Contract deleted successfully"}


# ================================
# Routes - Analytics & Stats
# ================================

@app.get("/api/stats")
async def get_stats(db: Session = Depends(get_db)):
    """Get contract statistics"""
    total_contracts = db.query(Contract).count()
    
    # Total revenue by currency
    revenue_by_currency = {}
    contracts = db.query(Contract).all()
    
    for contract in contracts:
        currency = contract.currency
        if currency not in revenue_by_currency:
            revenue_by_currency[currency] = 0
        revenue_by_currency[currency] += contract.fees_numeric
    
    # Most common services
    service_counts = {}
    for contract in contracts:
        for service in contract.services.split(", "):
            service_counts[service] = service_counts.get(service, 0) + 1
    
    return {
        "total_contracts": total_contracts,
        "revenue_by_currency": revenue_by_currency,
        "service_counts": service_counts,
        "recent_contracts": len([c for c in contracts if (datetime.utcnow() - c.created_at).days <= 30])
    }


# ================================
# Health Check
# ================================

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "version": "2.0",
        "timestamp": datetime.utcnow().isoformat()
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=4000)
